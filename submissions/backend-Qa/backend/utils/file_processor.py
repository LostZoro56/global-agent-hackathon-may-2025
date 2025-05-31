import os
import re
import mimetypes
from typing import Optional

# Try to import various document processing libraries
# These will need to be installed via pip
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Extract text content from various file types.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content or None if extraction failed
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # Extract text based on file type
    if ext == '.txt' or ext == '.feature':
        return extract_from_text_file(file_path)
    elif ext == '.docx':
        return extract_from_docx(file_path)
    elif ext == '.pdf':
        return extract_from_pdf(file_path)
    else:
        # Try to guess the file type based on content
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            if mime_type == 'text/plain':
                return extract_from_text_file(file_path)
            elif mime_type == 'application/pdf':
                return extract_from_pdf(file_path)
        
        # Default to treating as text file
        try:
            return extract_from_text_file(file_path)
        except:
            return None

def extract_from_text_file(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try different encodings if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except:
            raise ValueError(f"Could not decode text file: {file_path}")

def extract_from_docx(file_path: str) -> str:
    """Extract text from a .docx file."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx package is required for .docx file processing")
    
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)

def extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 package is required for PDF file processing")
    
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = []
        
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text.append(page.extract_text())
        
        return '\n'.join(text)
